from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import Response, StreamingResponse
from backend.app.db.models import Resume, JobMatch
from mcp.tools.file_tool import FileTool
from mcp.tools.llm_tool import LLMTool
from mcp.tools.browser_tool import BrowserTool
from backend.app.rag.retriever import rag_retriever
from typing import Dict, Optional
import io
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from pydantic import BaseModel

router = APIRouter(prefix="/resume", tags=["Resume"])

file_tool = FileTool()
llm_brain = LLMTool()
browser_tool = BrowserTool()

class AnalyzeRequest(BaseModel):
    resume_filename: str
    job_url: Optional[str] = None
    job_description: Optional[str] = None

class DownloadRequest(BaseModel):
    content: str
    filename: str

@router.post("/upload")
async def upload_resume(file: UploadFile = File(...)):
    # 1. Save file to disk
    file_path = f"data/resumes/{file.filename}"
    file_bytes = await file.read()
    await file_tool.execute("save_upload", file_path, file_bytes)
    
    # 2. Extract text from the saved file
    content = await file_tool.execute("read_pdf", file_path)
    if not content or str(content).startswith("File not found"):
        raise HTTPException(status_code=500, detail=f"Failed to parse resume after saving: {content}")
    
    # 3. Index in RAG
    rag_retriever.add_resume(file.filename, content, {"filename": file.filename})
    
    return {"message": "Resume uploaded and indexed", "filename": file.filename}

@router.post("/analyze")
async def analyze_job(request: AnalyzeRequest):
    # 1. Get JD (either from URL or direct text)
    jd = ""
    job_title = "Tailored Job"
    
    if request.job_url:
        job_data = await browser_tool.execute(request.job_url)
        jd = job_data["content"]
        job_title = job_data["title"]
    elif request.job_description:
        jd = request.job_description
        # Try to extract a title from the first line or first 50 chars
        job_title = jd.split('\n')[0][:50] if jd else "Tailored Job"
    else:
        raise HTTPException(status_code=400, detail="Provide either job_url or job_description")
    
    # 2. Get resume content from RAG
    resume_content = rag_retriever.get_resume(request.resume_filename)
    if not resume_content:
        raise HTTPException(status_code=404, detail=f"Resume '{request.resume_filename}' not found in AI vector database")
    
    # 3. Brain analysis
    brain_res = await llm_brain.execute("tailor_resume", {
        "resume": resume_content,
        "job_description": jd
    })
    
    return {
        "job_title": job_title,
        "original_resume": resume_content,
        "tailored_resume": brain_res["tailored_resume"],
        "ats_score": brain_res["ats_score"]
    }

@router.post("/download/pdf")
async def download_pdf(request: DownloadRequest):
    buffer = io.BytesIO()
    
    # Establish document template
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    
    # Custom styles for more "premium" look
    styles.add(ParagraphStyle(
        name='ResumeHeader',
        parent=styles['Heading1'],
        alignment=TA_CENTER,
        fontSize=18,
        spaceAfter=12
    ))
    
    styles.add(ParagraphStyle(
        name='ResumeBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=6
    ))

    content = []
    
    # Process text into paragraphs
    text_blocks = request.content.split('\n\n')
    for block in text_blocks:
        if not block.strip():
            continue
            
        # Try to detect headers (simple heuristic)
        if block.isupper() or len(block.split('\n')[0]) < 30 and block.split('\n')[0].isupper():
            content.append(Paragraph(block.replace('\n', '<br/>'), styles['ResumeHeader']))
        else:
            content.append(Paragraph(block.replace('\n', '<br/>'), styles['ResumeBody']))
        
        content.append(Spacer(1, 0.1 * inch))

    # Build PDF
    doc.build(content)
    
    buffer.seek(0)
    # Sanitize filename
    import re
    safe_filename = re.sub(r'[^a-zA-Z0-9_\-]', '_', request.filename)
    
    return StreamingResponse(
        buffer, 
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={safe_filename}.pdf"}
    )

@router.post("/download/docx")
async def download_docx(request: DownloadRequest):
    doc = Document()
    
    # Process text into paragraphs
    text_blocks = request.content.split('\n\n')
    for block in text_blocks:
        if not block.strip():
            continue
            
        # Detect headers for bolding
        if block.isupper() or len(block.split('\n')[0]) < 30 and block.split('\n')[0].isupper():
            p = doc.add_paragraph()
            run = p.add_run(block)
            run.bold = True
        else:
            doc.add_paragraph(block)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Sanitize filename
    import re
    safe_filename = re.sub(r'[^a-zA-Z0-9_\-]', '_', request.filename)
    
    return StreamingResponse(
        buffer, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={safe_filename}.docx"}
    )
